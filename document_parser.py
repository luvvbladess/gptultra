"""
Модуль для извлечения текста из документов (DOCX, PDF)
"""

import io
from typing import Optional

from docx import Document
import fitz  # PyMuPDF


async def extract_text_from_docx(file_data: bytes) -> str:
    """
    Извлекает текст из DOCX файла.
    
    Args:
        file_data: Байты файла DOCX
        
    Returns:
        Извлеченный текст
    """
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Также извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    except Exception as e:
        return f"Ошибка при чтении DOCX: {str(e)}"


async def extract_text_from_pdf(file_data: bytes) -> str:
    """
    Извлекает текст из PDF файла.
    
    Args:
        file_data: Байты файла PDF
        
    Returns:
        Извлеченный текст
    """
    try:
        text_parts = []
        
        # Открываем PDF из байтов
        pdf_document = fitz.open(stream=file_data, filetype="pdf")
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            if text.strip():
                text_parts.append(f"--- Страница {page_num + 1} ---\n{text}")
        
        pdf_document.close()
        
        return "\n\n".join(text_parts)
    except Exception as e:
        return f"Ошибка при чтении PDF: {str(e)}"


async def extract_text_from_txt(file_data: bytes) -> str:
    """
    Извлекает текст из TXT файла с автоопределением кодировки.
    
    Args:
        file_data: Байты файла TXT
        
    Returns:
        Извлеченный текст
    """
    # Пробуем разные кодировки
    encodings = ['utf-8', 'utf-8-sig', 'cp1251', 'cp1252', 'latin-1', 'iso-8859-1', 'koi8-r']
    
    for encoding in encodings:
        try:
            return file_data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    
    # Если ничего не подошло, декодируем с ошибками
    return file_data.decode('utf-8', errors='replace')


async def extract_text_from_file(file_data: bytes, file_name: str) -> Optional[str]:
    """
    Определяет тип файла и извлекает текст.
    
    Args:
        file_data: Байты файла
        file_name: Имя файла
        
    Returns:
        Извлеченный текст или None, если формат не поддерживается
    """
    file_name_lower = file_name.lower()
    
    if file_name_lower.endswith('.docx'):
        return await extract_text_from_docx(file_data)
    elif file_name_lower.endswith('.pdf'):
        return await extract_text_from_pdf(file_data)
    elif file_name_lower.endswith('.txt'):
        return await extract_text_from_txt(file_data)
    else:
        return None


async def edit_docx_with_replacements(file_data: bytes, replacements: dict) -> bytes:
    """
    Редактирует DOCX документ, заменяя текст с сохранением форматирования.
    
    Args:
        file_data: Байты исходного DOCX файла
        replacements: Словарь замен {старый_текст: новый_текст}
        
    Returns:
        Байты отредактированного документа
    """
    doc = Document(io.BytesIO(file_data))
    
    def replace_in_paragraph(paragraph, old_text, new_text):
        """Заменяет текст в параграфе, сохраняя форматирование первого run"""
        if old_text not in paragraph.text:
            return False
        
        # Собираем полный текст и находим позицию замены
        full_text = paragraph.text
        
        if old_text in full_text:
            # Простой случай - текст находится целиком в одном run
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    return True
            
            # Сложный случай - текст разбит по нескольким runs
            # Объединяем все runs в один с сохранением форматирования первого
            if paragraph.runs:
                first_run = paragraph.runs[0]
                new_full_text = full_text.replace(old_text, new_text)
                
                # Очищаем все runs кроме первого
                for run in paragraph.runs[1:]:
                    run.text = ""
                
                first_run.text = new_full_text
                return True
        
        return False
    
    def replace_in_cell(cell, old_text, new_text):
        """Заменяет текст в ячейке таблицы"""
        replaced = False
        for paragraph in cell.paragraphs:
            if replace_in_paragraph(paragraph, old_text, new_text):
                replaced = True
        return replaced
    
    # Заменяем во всех параграфах документа
    for old_text, new_text in replacements.items():
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, old_text, new_text)
        
        # Заменяем в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_cell(cell, old_text, new_text)
        
        # Заменяем в headers и footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    replace_in_paragraph(paragraph, old_text, new_text)
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    replace_in_paragraph(paragraph, old_text, new_text)
    
    # Сохраняем в байты
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


async def get_docx_structure_for_ai(file_data: bytes) -> str:
    """
    Извлекает структуру документа для анализа AI.
    Возвращает текст с маркерами для понимания структуры.
    """
    doc = Document(io.BytesIO(file_data))
    parts = []
    
    parts.append("=== СОДЕРЖИМОЕ ДОКУМЕНТА ===\n")
    
    # Параграфы
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            style_name = paragraph.style.name if paragraph.style else "Normal"
            parts.append(f"[П{i+1}|{style_name}] {paragraph.text}")
    
    # Таблицы
    for t_idx, table in enumerate(doc.tables):
        parts.append(f"\n--- ТАБЛИЦА {t_idx + 1} ---")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(f"  Строка {r_idx + 1}: {' | '.join(cells)}")
    
    return "\n".join(parts)
