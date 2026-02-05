"""
Модуль для генерации DOCX документов из Markdown/текста.
Использует связку markdown -> html -> docx для поддержки таблиц и форматирования.
"""

import io
import markdown
from typing import Optional
import markdown
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from htmldocx import HtmlToDocx

def create_list_numbering(doc, abstract_num_id=1):
    """
    Creates a new numbering definition (w:num) that points to the given abstractNumId.
    Returns the new numId.
    """
    # Get the numbering part
    try:
        numbering_part = doc.part.numbering_part
    except NotImplementedError:
        # Happens in some simplified docx structures or testing
        return None
        
    # Generate a new numId
    # Simply finding the max and adding 1 is a safe strategy usually, 
    # but python-docx doesn't expose a clean add_num method that returns ID easily.
    # We'll use the internal _next_numId if available or manual calculation.
    
    # Access the numbering element
    numbering_element = numbering_part.numbering_definitions._numbering
    
    # Find next available ID
    current_ids = [int(num.numId) for num in numbering_element.num_lst]
    next_num_id = max(current_ids, default=0) + 1
    
    # Create new <w:num w:numId="X">
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(next_num_id))
    
    # <w:abstractNumId w:val="Y"/>
    abstractNumId = OxmlElement('w:abstractNumId')
    abstractNumId.set(qn('w:val'), str(abstract_num_id))
    num.append(abstractNumId)
    
    # Force restart at 1 using lvlOverride
    # This is the "nuclear option" for Word numbering that usually works even if abstractNum is reused
    lvlOverride = OxmlElement('w:lvlOverride')
    lvlOverride.set(qn('w:ilvl'), '0')
    startOverride = OxmlElement('w:startOverride')
    startOverride.set(qn('w:val'), '1')
    lvlOverride.append(startOverride)
    num.append(lvlOverride)

    
    # Add to numbering part
    numbering_element.append(num)
    
    return next_num_id

def fix_numbered_lists(doc):
    """
    Post-process the document to reset numbering for separate lists.
    htmldocx typically uses 'List Number' style for <ol>.
    We detect breaks in list paragraphs (including tables!) and force a new numId.
    """
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.text.paragraph import Paragraph

    last_was_list_number = False
    current_num_id_val = None
    
    list_style_name = 'List Number' # Default in htmldocx for <ol>
    abstract_num_id = None
    
    # Try to find the abstractNumId used by 'List Number' style
    try:
        styles = doc.styles
        style_options = [list_style_name, 'List Paragraph']
        for s_name in style_options:
            if s_name in styles:
                style = styles[s_name]

            if hasattr(style, '_element') and style._element.pPr is not None and style._element.pPr.numPr is not None:
                default_num_id = style._element.pPr.numPr.numId.val
                # Now finding abstractNumId from numbering part
                numbering_part = doc.part.numbering_part
                num = numbering_part.numbering_definitions._numbering.get_num(default_num_id)
                if num is not None:
                    abstract_num_id = num.abstractNumId.val
            if abstract_num_id is not None:
                break
    except Exception as e:
        pass
        
    # If we couldn't find it, we'll try to guess or find the first one
    if abstract_num_id is None:
        try:
             # Fallback: assume 1
             abstract_num_id = 1
        except:
             return

    # Iterate over ALL elements in the body (paragraphs AND tables)
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            
            # Check if this is a List Number paragraph
            if p.style.name == list_style_name or p.style.name == 'List Paragraph':

                # If previous element was NOT List Number, this is a START of a new list
                if not last_was_list_number:
                    # We need a NEW numbering ID
                    current_num_id_val = create_list_numbering(doc, abstract_num_id)
                
                # Apply the current_num_id_val to this paragraph if we have one
                if current_num_id_val is not None:
                    pPr = p._element.get_or_add_pPr()
                    numPr = pPr.get_or_add_numPr()
                    # Use get_or_add_numId() - this returns a CT_DecimalNumber
                    numId_element = numPr.get_or_add_numId()
                    numId_element.set(qn('w:val'), str(current_num_id_val))
                    
                    # Ensure level is 0
                    ilvl = numPr.get_or_add_ilvl()
                    ilvl.set(qn('w:val'), '0') 
                
                last_was_list_number = True
            else:
                # Normal paragraph -> Break
                last_was_list_number = False
        
        elif isinstance(child, CT_Tbl):
            # Table -> Break the list sequence
            last_was_list_number = False
        
        else:
            # Any other element -> Break
            last_was_list_number = False



def convert_markdown_to_docx(markdown_text: str) -> bytes:
    """
    Конвертирует Markdown текст в DOCX документ.
    
    Args:
        markdown_text: Исходный текст в формате Markdown
        
    Returns:
        Байты сгенерированного DOCX файла
    """
    # 0. Препроцессинг текста
    import re
    
    lines = markdown_text.split('\n')
    fixed_lines = []
    
    for i, line in enumerate(lines):
        # 1. Нормализация нумерации списков: заменяем "12. " на "1. "
        # Это позволяет htmldocx/word корректно обрабатывать их как списки, а наша пост-обработка сбросит нумерацию где надо.
        # Если оставить "12.", markdown сделает <ol start="12">, и Word может это криво понять (или мы не сможем сбросить).
        # Заменяем только если это начало строки (с учетом отступов)
        line = re.sub(r'^(\s*)\d+\.', r'\1 1.', line)
        
        # 2. Исправление таблиц
        # Если строка похожа на начало таблицы (содержит | и не пустая), а предыдущая не пустая - добавляем отступ
        if "|" in line and i > 0 and lines[i-1].strip() and not lines[i-1].strip().startswith("|"):
             # Проверяем, что это действительно таблица (наличие разделителя ---|--- на следующей строке)
             if i + 1 < len(lines) and set(lines[i+1].strip()) <= set("|-:| "):
                 fixed_lines.append("")
        
        fixed_lines.append(line)
        
    markdown_text = "\n".join(fixed_lines)

    # 1. Конвертируем Markdown в HTML
    # Используем расширения для поддержки таблиц и других элементов
    html_text = markdown.markdown(
        markdown_text, 
        extensions=['tables', 'extra', 'fenced_code', 'nl2br']
    )
    
    # 2. Создаем документ и парсер
    doc = Document()
    new_parser = HtmlToDocx()
    
    # 3. Парсим HTML и добавляем в документ
    # Оборачиваем в try-except на случай проблем с парсингом сложного HTML
    try:
        new_parser.add_html_to_document(html_text, doc)
    except Exception as e:
        # В случае ошибки добавляем текст как есть, но сигнализируем об ошибке в лог (или просто добавляем параграф)
        print(f"HTMLDOCX CRITICAL ERROR: {e}")
        import traceback
        traceback.print_exc()
        doc.add_paragraph("Ошибка при конвертации форматирования. Исходный текст:")
        doc.add_paragraph(markdown_text)
    
    # 5. Исправляем нумерацию списков (сброс нумерации для новых списков)
    try:
        fix_numbered_lists(doc)
    except Exception as e:
        # Не ломаем генерацию, если фикс не сработал
        print(f"List fix warning: {e}")

    # 6. Сохраняем в байты
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.read()
