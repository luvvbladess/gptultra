"""
Обработчики сообщений Telegram бота
"""

import logging
import asyncio
from typing import Optional

from aiogram import Router, Bot, F
from aiogram.types import Message, CallbackQuery, BufferedInputFile
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
import uuid
from io import BytesIO
from docx import Document

from config import (
    SYSTEM_PROMPT, 
    MAX_HISTORY_MESSAGES,
    MAX_TELEGRAM_MESSAGE_LENGTH,
    AVAILABLE_MODELS
)
from aiogram.types import ReactionTypeEmoji
from openai_client import get_chat_response, encode_image_to_base64, generate_image, edit_image_with_dalle, transcribe_audio
from document_parser import extract_text_from_file, edit_docx_with_replacements, get_docx_structure_for_ai
from conversations import conversation_manager
from keyboards import (
    get_main_menu_keyboard,
    get_conversations_keyboard,
    get_conversation_actions_keyboard,
    get_confirm_delete_keyboard,
    get_confirm_clear_keyboard,
    get_cancel_keyboard,
    get_models_keyboard,
    get_models_keyboard,
    get_custom_prompts_keyboard,
    get_txt_download_keyboard
)

# Глобальный кэш для хранения длинных ответов
# Key: UUID string, Value: text content
RESPONSE_CACHE = {}


logger = logging.getLogger(__name__)
router = Router()


class AnimatedLoader:
    """Анимированный лоадер с редактированием сообщения"""
    
    def __init__(self, message: Message, base_text: str = "Ищу ответ на Ваш вопрос"):
        self.message = message
        self.base_text = base_text
        self.status_msg = None
        self.running = False
        self.task = None
    
    async def start(self) -> Message:
        """Запускает анимацию"""
        self.status_msg = await self.message.reply(f"{self.base_text}...")
        self.running = True
        self.task = asyncio.create_task(self._animate())
        return self.status_msg
    
    async def _animate(self):
        """Анимация точек"""
        dots = [".", "..", "..."]
        i = 0
        while self.running:
            try:
                await asyncio.sleep(0.8)
                if not self.running:
                    break
                i = (i + 1) % len(dots)
                await self.status_msg.edit_text(f"{self.base_text}{dots[i]}")
            except Exception:
                break
    
    async def stop_with_result(self, result_text: str = None):
        """Останавливает анимацию и заменяет на результат"""
        self.running = False
        if self.task:
            self.task.cancel()
            try:
                await self.task
            except asyncio.CancelledError:
                pass
        return self.status_msg
    
    async def stop(self):
        """Просто останавливает анимацию"""
        self.running = False
        if self.task:
            self.task.cancel()


# Состояния FSM
def safe_handler(func):
    """Декоратор для безопасного выполнения хендлеров"""
    async def wrapper(message: Message, bot: Bot, *args, **kwargs):
        try:
            return await func(message, bot, *args, **kwargs)
        except Exception as e:
            logger.error(f"Unhandled error in {func.__name__}: {e}")
            try:
                await message.reply(f"❌ Критическая ошибка: {e}")
            except Exception:
                pass
    return wrapper


class BotStates(StatesGroup):
    waiting_for_rename = State()
    waiting_for_custom_prompt = State()


async def add_heart_reaction(message: Message, bot: Bot) -> None:
    """Добавляет реакцию ❤️ на сообщение пользователя"""
    try:
        await bot.set_message_reaction(
            chat_id=message.chat.id,
            message_id=message.message_id,
            reaction=[ReactionTypeEmoji(emoji="❤")]
        )
    except Exception as e:
        # Игнорируем ошибки реакций (могут быть отключены в чате)
        logger.debug(f"Could not add reaction: {e}")


def get_updated_keyboard(user_id: int) -> None:
    """Возвращает клавиатуру с учётом текущих режимов пользователя"""
    is_dalle = conversation_manager.is_dalle_mode(user_id)
    is_edit = conversation_manager.is_edit_mode(user_id)
    is_template = conversation_manager.is_template_mode(user_id)
    return get_main_menu_keyboard(is_dalle_mode=is_dalle, is_edit_mode=is_edit, is_template_mode=is_template)


@router.message(Command("start"))
async def cmd_start(message: Message) -> None:
    """Обработчик команды /start"""
    user_id = message.from_user.id
    
    # Создаём первую беседу, если её нет
    if not conversation_manager.get_conversations(user_id):
        conversation_manager.create_conversation(user_id, "Новая беседа")
    
    current_model = conversation_manager.get_user_model(user_id)
    model_name = AVAILABLE_MODELS.get(current_model, current_model)
    
    welcome_text = f"""👋 **Привет!**

Я — AI-ассистент на базе OpenAI GPT.

**Что я умею:**
• 💬 Вести несколько бесед параллельно
• 🖼 Анализировать изображения
• 📄 Читать PDF и DOCX документы
• 🤖 Переключаться между моделями
• ✨ Кастомные промпты (до 2 шт.)

**Текущая модель:** {model_name}

Используй меню ниже для управления.
Просто напиши сообщение, чтобы начать! 🚀"""
    
    await message.answer(welcome_text, reply_markup=get_updated_keyboard(user_id), parse_mode="Markdown")


@router.message(Command("help"))
async def cmd_help(message: Message) -> None:
    """Обработчик команды /help"""
    await show_help(message)


# ============== КНОПКИ МЕНЮ ==============

@router.message(F.text == "📝 Новая беседа")
async def btn_new_conversation(message: Message) -> None:
    """Создание новой беседы"""
    user_id = message.from_user.id
    conv = conversation_manager.create_conversation(user_id)
    
    await message.answer(
        f"✅ Создана новая беседа: **{conv.title}**\n\n"
        "Напиши мне что-нибудь!",
        parse_mode="Markdown"
    )


@router.message(F.text == "📂 Мои беседы")
async def btn_my_conversations(message: Message) -> None:
    """Показать список бесед"""
    user_id = message.from_user.id
    
    conversations = conversation_manager.get_conversations(user_id)
    active = conversation_manager.get_active_conversation(user_id)
    active_id = active.id if active else None
    
    if not conversations:
        text = "У тебя пока нет бесед. Создай первую!"
    else:
        text = f"📂 **Твои беседы** ({len(conversations)}):\n\n"
        text += "Выбери беседу или создай новую:"
    
    await message.answer(
        text,
        reply_markup=get_conversations_keyboard(conversations, active_id),
        parse_mode="Markdown"
    )


@router.message(F.text == "🤖 Модель")
async def btn_select_model(message: Message) -> None:
    """Показать выбор модели"""
    user_id = message.from_user.id
    current_model = conversation_manager.get_user_model(user_id)
    
    await message.answer(
        "🤖 **Выбери модель:**\n\n"
        "Разные модели имеют разные возможности и скорость.",
        reply_markup=get_models_keyboard(current_model),
        parse_mode="Markdown"
    )


@router.message(F.text.in_(["🎨 Редактор", "❌ Выйти из редактора"]))
async def btn_editor_mode(message: Message) -> None:
    """Включить/выключить режим редактирования изображений"""
    user_id = message.from_user.id
    
    is_edit_mode = conversation_manager.is_edit_mode(user_id)
    
    if is_edit_mode:
        # Выключаем режим
        conversation_manager.set_edit_mode(user_id, False)
        await message.answer(
            "🎨 **Режим редактирования выключен**\n\n"
            "Теперь бот работает в обычном режиме.",
            parse_mode="Markdown",
            reply_markup=get_updated_keyboard(user_id)
        )
    else:
        # Включаем режим
        conversation_manager.set_edit_mode(user_id, True)
        await message.answer(
            "🎨 **Режим редактирования включён!**\n\n"
            "📸 **Как использовать:**\n"
            "1. Отправь фото, которое хочешь редактировать\n"
            "2. Напиши что изменить (например: \"добавь шляпу\")\n"
            "3. Бот запомнит картинку и можно продолжать редактировать\n\n"
            "💡 Картинка сохраняется до выхода из режима.\n"
            "Нажми ❌ Выйти из редактора чтобы выйти.",
            parse_mode="Markdown",
            reply_markup=get_updated_keyboard(user_id)
        )


@router.message(F.text.in_(["🖼 DALL-E", "❌ Выйти из DALL-E"]))
async def btn_dalle_mode(message: Message) -> None:
    """Включить/выключить DALL-E режим генерации"""
    user_id = message.from_user.id
    
    is_dalle_mode = conversation_manager.is_dalle_mode(user_id)
    
    if is_dalle_mode:
        # Выключаем режим
        conversation_manager.set_dalle_mode(user_id, False)
        await message.answer(
            "🖼 **DALL-E режим выключен**\n\n"
            "Теперь бот работает в обычном режиме.",
            parse_mode="Markdown",
            reply_markup=get_updated_keyboard(user_id)
        )
    else:
        # Включаем режим
        conversation_manager.set_dalle_mode(user_id, True)
        await message.answer(
            "🖼 **DALL-E режим включён!**\n\n"
            "🎨 **Как использовать:**\n"
            "• Просто напиши что нарисовать\n"
            "• Бот запомнит последнюю картинку\n"
            "• Можешь дописать: \"добавь солнце\" - и он отредактирует\n\n"
            "💡 Первое сообщение = новая картинка\n"
            "Следующие = редактирование\n\n"
            "Нажми ❌ Выйти из DALL-E чтобы выйти.",
            parse_mode="Markdown",
            reply_markup=get_updated_keyboard(user_id)
        )


@router.message(F.text.in_(["📄 Шаблоны", "❌ Выйти из шаблона"]))
async def btn_template_mode(message: Message) -> None:
    """Включить/выключить режим шаблонов документов"""
    user_id = message.from_user.id
    
    is_template_mode = conversation_manager.is_template_mode(user_id)
    
    if is_template_mode:
        # Выключаем режим
        conversation_manager.set_template_mode(user_id, False)
        await message.answer(
            "📄 **Режим шаблонов выключен**\n\n"
            "Загруженный шаблон удалён. Теперь бот работает в обычном режиме.",
            parse_mode="Markdown",
            reply_markup=get_updated_keyboard(user_id)
        )
    else:
        # Включаем режим
        conversation_manager.set_template_mode(user_id, True)
        await message.answer(
            "📄 **Режим шаблонов включён!**\n\n"
            "📋 **Как использовать:**\n"
            "1. Отправь DOCX документ (шаблон или договор)\n"
            "2. Опиши какие изменения нужно сделать:\n"
            "   • _\"Замени ООО Ромашка на ООО Василёк\"_\n"
            "   • _\"Измени дату на 15.03.2026\"_\n"
            "   • _\"Поменяй сумму 100000 на 250000\"_\n"
            "3. Бот создаст новый документ с сохранением форматирования!\n\n"
            "💡 Шрифты, стили и оформление сохранятся.\n"
            "Нажми ❌ Выйти из шаблона чтобы выйти.",
            parse_mode="Markdown",
            reply_markup=get_updated_keyboard(user_id)
        )


@router.message(F.text == "🗑 Очистить")
async def btn_clear_current(message: Message) -> None:
    """Очистка текущей беседы"""
    user_id = message.from_user.id
    conv = conversation_manager.get_active_conversation(user_id)
    
    if not conv:
        await message.answer("❌ Нет активной беседы.")
        return
    
    await message.answer(
        f"🧹 Очистить историю беседы **{conv.title}**?\n\n"
        "Все сообщения будут удалены.",
        reply_markup=get_confirm_clear_keyboard(conv.id),
        parse_mode="Markdown"
    )


@router.message(F.text == "ℹ️ Помощь")
async def btn_help(message: Message) -> None:
    """Показать помощь"""
    await show_help(message)


async def show_help(message: Message) -> None:
    """Выводит справку"""
    user_id = message.from_user.id
    current_model = conversation_manager.get_user_model(user_id)
    model_name = AVAILABLE_MODELS.get(current_model, current_model)
    
    help_text = f"""📖 **Справка**

**Команды:**
• `/start` — начать работу
• `/help` — показать справку

**Меню:**
• 📝 **Новая беседа** — создать новую беседу
• 📂 **Мои беседы** — список всех бесед
• 🤖 **Модель** — сменить модель AI
• 🗑 **Очистить** — очистить историю
• ℹ️ **Помощь** — эта справка

**Доступные модели:**
• ⚡ GPT-5 Nano — быстрая
• 🔹 GPT-5 Mini — баланс
• 🔷 GPT-5.2 — умная
• 💎 GPT-5.2 Pro — максимум

**Текущая модель:** {model_name}

**Возможности:**
• Веду несколько бесед с отдельной историей
• Анализирую фото (отправь картинку)
• Читаю PDF и DOCX файлы"""
    
    await message.answer(help_text, parse_mode="Markdown")


@router.message(F.text == "✨ Промпты")
async def btn_custom_prompts(message: Message) -> None:
    """Показать меню управления кастомными промптами"""
    user_id = message.from_user.id
    
    prompts = conversation_manager.get_custom_prompts(user_id)
    active = conversation_manager.get_active_custom_prompt(user_id)
    
    text = """✨ **Кастомные промпты**

Здесь ты можешь добавить до 2 своих системных промптов.
Кастомный промпт будет добавлен к стандартному.

"""
    if prompts:
        text += f"У тебя {len(prompts)} промпт(ов):\n\n"
        for i, p in enumerate(prompts):
            is_active = p == active
            status = "✅ активен" if is_active else ""
            text += f"**{i+1}.** {p[:50]}{'...' if len(p) > 50 else ''} {status}\n\n"
    else:
        text += "_Ещё нет сохранённых промптов._\n\n"
    
    if active:
        text += "🟢 Сейчас используется кастомный промпт."
    else:
        text += "⚪ Сейчас используется стандартный промпт."
    
    await message.answer(
        text,
        reply_markup=get_custom_prompts_keyboard(prompts, active),
        parse_mode="Markdown"
    )


# ============== CALLBACK HANDLERS ==============

@router.callback_query(F.data == "new_conversation")
async def callback_new_conversation(callback: CallbackQuery) -> None:
    """Создание новой беседы через inline кнопку"""
    user_id = callback.from_user.id
    conv = conversation_manager.create_conversation(user_id)
    
    await callback.message.edit_text(
        f"✅ Создана новая беседа: **{conv.title}**\n\n"
        "Напиши мне что-нибудь!",
        parse_mode="Markdown"
    )
    await callback.answer()


@router.callback_query(F.data == "list_conversations")
async def callback_list_conversations(callback: CallbackQuery) -> None:
    """Показать список бесед"""
    user_id = callback.from_user.id
    
    conversations = conversation_manager.get_conversations(user_id)
    active = conversation_manager.get_active_conversation(user_id)
    active_id = active.id if active else None
    
    text = f"📂 **Твои беседы** ({len(conversations)}):\n\n"
    text += "Выбери беседу или создай новую:"
    
    await callback.message.edit_text(
        text,
        reply_markup=get_conversations_keyboard(conversations, active_id),
        parse_mode="Markdown"
    )
    await callback.answer()


@router.callback_query(F.data.startswith("select_conv:"))
async def callback_select_conversation(callback: CallbackQuery) -> None:
    """Выбор беседы"""
    user_id = callback.from_user.id
    conv_id = callback.data.split(":")[1]
    
    conv = conversation_manager.set_active_conversation(user_id, conv_id)
    
    if conv:
        await callback.message.edit_text(
            f"📝 **{conv.title}**\n\n"
            f"Сообщений в истории: {len(conv.messages)}\n\n"
            "Что хочешь сделать?",
            reply_markup=get_conversation_actions_keyboard(conv_id),
            parse_mode="Markdown"
        )
    else:
        await callback.answer("❌ Беседа не найдена", show_alert=True)
    
    await callback.answer()


@router.callback_query(F.data.startswith("select_model:"))
async def callback_select_model(callback: CallbackQuery) -> None:
    """Выбор модели"""
    user_id = callback.from_user.id
    model_id = callback.data.split(":")[1]
    
    if model_id in AVAILABLE_MODELS:
        conversation_manager.set_user_model(user_id, model_id)
        model_name = AVAILABLE_MODELS[model_id]
        
        await callback.message.edit_text(
            f"✅ Модель изменена на: **{model_name}**\n\n"
            "Новые сообщения будут обрабатываться этой моделью.",
            parse_mode="Markdown"
        )
        await callback.answer(f"Выбрана модель: {model_id}")
    else:
        await callback.answer("❌ Модель не найдена", show_alert=True)


@router.callback_query(F.data.startswith("rename_conv:"))
async def callback_rename_conversation(callback: CallbackQuery, state: FSMContext) -> None:
    """Начать переименование беседы"""
    conv_id = callback.data.split(":")[1]
    
    await state.set_state(BotStates.waiting_for_rename)
    await state.update_data(conv_id=conv_id)
    
    await callback.message.edit_text(
        "✏️ Введи новое название для беседы:",
        reply_markup=get_cancel_keyboard()
    )
    await callback.answer()


@router.callback_query(F.data == "cancel_action")
async def callback_cancel_action(callback: CallbackQuery, state: FSMContext) -> None:
    """Отмена действия"""
    await state.clear()
    
    user_id = callback.from_user.id
    conversations = conversation_manager.get_conversations(user_id)
    active = conversation_manager.get_active_conversation(user_id)
    active_id = active.id if active else None
    
    await callback.message.edit_text(
        "❌ Действие отменено.\n\n📂 Твои беседы:",
        reply_markup=get_conversations_keyboard(conversations, active_id)
    )
    await callback.answer()


@router.message(BotStates.waiting_for_rename)
async def process_rename(message: Message, state: FSMContext) -> None:
    """Обработка нового названия беседы"""
    user_id = message.from_user.id
    data = await state.get_data()
    conv_id = data.get("conv_id")
    
    new_title = message.text.strip()[:50]  # Ограничиваем длину
    
    if conversation_manager.rename_conversation(user_id, conv_id, new_title):
        await message.answer(f"✅ Беседа переименована в: **{new_title}**", parse_mode="Markdown")
    else:
        await message.answer("❌ Не удалось переименовать беседу.")
    
    await state.clear()


@router.callback_query(F.data.startswith("clear_conv:"))
async def callback_clear_conversation(callback: CallbackQuery) -> None:
    """Запрос на очистку беседы"""
    conv_id = callback.data.split(":")[1]
    
    await callback.message.edit_text(
        "🧹 Очистить всю историю этой беседы?\n\n"
        "⚠️ Все сообщения будут удалены!",
        reply_markup=get_confirm_clear_keyboard(conv_id)
    )
    await callback.answer()


@router.callback_query(F.data.startswith("confirm_clear:"))
async def callback_confirm_clear(callback: CallbackQuery) -> None:
    """Подтверждение очистки беседы"""
    user_id = callback.from_user.id
    conv_id = callback.data.split(":")[1]
    
    if conversation_manager.clear_conversation(user_id, conv_id):
        await callback.message.edit_text("✅ История беседы очищена!")
    else:
        await callback.message.edit_text("❌ Не удалось очистить беседу.")
    
    await callback.answer()


@router.callback_query(F.data.startswith("delete_conv:"))
async def callback_delete_conversation(callback: CallbackQuery) -> None:
    """Запрос на удаление беседы"""
    conv_id = callback.data.split(":")[1]
    
    await callback.message.edit_text(
        "🗑 Удалить эту беседу?\n\n"
        "⚠️ Это действие нельзя отменить!",
        reply_markup=get_confirm_delete_keyboard(conv_id)
    )
    await callback.answer()


@router.callback_query(F.data.startswith("confirm_delete:"))
async def callback_confirm_delete(callback: CallbackQuery) -> None:
    """Подтверждение удаления беседы"""
    user_id = callback.from_user.id
    conv_id = callback.data.split(":")[1]
    
    if conversation_manager.delete_conversation(user_id, conv_id):
        # Создаём новую беседу если удалили последнюю
        if not conversation_manager.get_conversations(user_id):
            conversation_manager.create_conversation(user_id, "Новая беседа")
        
        await callback.message.edit_text("✅ Беседа удалена!")
    else:
        await callback.message.edit_text("❌ Не удалось удалить беседу.")
    
    await callback.answer()


# ============== КАСТОМНЫЕ ПРОМПТЫ CALLBACKS ==============

@router.callback_query(F.data.startswith("toggle_prompt:"))
async def callback_toggle_prompt(callback: CallbackQuery) -> None:
    """Включить/выключить кастомный промпт"""
    user_id = callback.from_user.id
    index = int(callback.data.split(":")[1])
    
    prompts = conversation_manager.get_custom_prompts(user_id)
    active = conversation_manager.get_active_custom_prompt(user_id)
    
    if 0 <= index < len(prompts):
        selected_prompt = prompts[index]
        
        if selected_prompt == active:
            # Отключаем, если уже активен
            conversation_manager.set_active_custom_prompt(user_id, None)
            await callback.answer("Кастомный промпт отключён")
        else:
            # Включаем выбранный
            conversation_manager.set_active_custom_prompt(user_id, index)
            await callback.answer(f"Промпт {index + 1} активирован!")
        
        # Обновляем сообщение
        prompts = conversation_manager.get_custom_prompts(user_id)
        active = conversation_manager.get_active_custom_prompt(user_id)
        
        await callback.message.edit_reply_markup(
            reply_markup=get_custom_prompts_keyboard(prompts, active)
        )
    else:
        await callback.answer("❌ Промпт не найден", show_alert=True)


@router.callback_query(F.data.startswith("delete_prompt:"))
async def callback_delete_prompt(callback: CallbackQuery) -> None:
    """Удалить кастомный промпт"""
    user_id = callback.from_user.id
    index = int(callback.data.split(":")[1])
    
    if conversation_manager.delete_custom_prompt(user_id, index):
        await callback.answer(f"Промпт {index + 1} удалён!")
        
        # Обновляем сообщение
        prompts = conversation_manager.get_custom_prompts(user_id)
        active = conversation_manager.get_active_custom_prompt(user_id)
        
        text = """✨ **Кастомные промпты**

Здесь ты можешь добавить до 2 своих системных промптов.

"""
        if prompts:
            text += f"У тебя {len(prompts)} промпт(ов).\n"
        else:
            text += "_Все промпты удалены._\n"
        
        await callback.message.edit_text(
            text,
            reply_markup=get_custom_prompts_keyboard(prompts, active),
            parse_mode="Markdown"
        )
    else:
        await callback.answer("❌ Не удалось удалить промпт", show_alert=True)


@router.callback_query(F.data == "add_custom_prompt")
async def callback_add_custom_prompt(callback: CallbackQuery, state: FSMContext) -> None:
    """Начать добавление кастомного промпта"""
    user_id = callback.from_user.id
    prompts = conversation_manager.get_custom_prompts(user_id)
    
    note = ""
    if len(prompts) >= 2:
        note = "\n\n⚠️ У тебя уже 2 промпта. Новый заменит самый старый."
    
    await state.set_state(BotStates.waiting_for_custom_prompt)
    
    await callback.message.edit_text(
        f"✏️ **Введи текст нового промпта:**\n\n"
        "Этот текст будет добавлен к стандартному системному промпту.\n\n"
        "Например:\n"
        "• _Отвечай кратко, не более 3 предложений_\n"
        "• _Всегда предлагай примеры кода_\n"
        "• _Говори как пират_ 🏴‍☠️"
        f"{note}",
        reply_markup=get_cancel_keyboard(),
        parse_mode="Markdown"
    )
    await callback.answer()


@router.callback_query(F.data == "disable_custom_prompt")
async def callback_disable_custom_prompt(callback: CallbackQuery) -> None:
    """Отключить кастомный промпт"""
    user_id = callback.from_user.id
    
    conversation_manager.set_active_custom_prompt(user_id, None)
    
    prompts = conversation_manager.get_custom_prompts(user_id)
    
    await callback.message.edit_text(
        "✅ Теперь используется стандартный промпт.\n\n"
        "Твои сохранённые промпты по-прежнему доступны.",
        reply_markup=get_custom_prompts_keyboard(prompts, None),
        parse_mode="Markdown"
    )
    await callback.answer("Стандартный промпт активирован")


@router.callback_query(F.data == "no_action")
async def callback_no_action(callback: CallbackQuery) -> None:
    """Пустое действие для неактивных кнопок"""
    await callback.answer()


@router.message(BotStates.waiting_for_custom_prompt)
async def process_custom_prompt(message: Message, state: FSMContext) -> None:
    """Обработка нового кастомного промпта"""
    user_id = message.from_user.id
    prompt_text = message.text.strip()[:500]  # Ограничиваем длину
    
    if len(prompt_text) < 5:
        await message.answer("❌ Промпт слишком короткий. Минимум 5 символов.")
        return
    
    index = conversation_manager.add_custom_prompt(user_id, prompt_text)
    conversation_manager.set_active_custom_prompt(user_id, index - 1)  # Активируем новый
    
    await message.answer(
        f"✅ Промпт #{index} добавлен и активирован!\n\n"
        f"📝 _{prompt_text[:100]}{'...' if len(prompt_text) > 100 else ''}_\n\n"
        "Теперь бот будет использовать этот промпт.",
        parse_mode="Markdown"
    )
    
    await state.clear()


    
    await state.clear()


@router.callback_query(F.data.startswith("dl:"))
async def callback_download_response(callback: CallbackQuery) -> None:
    """Обработка скачивания TXT версии"""
    try:
        _, format_type, response_id = callback.data.split(":")
        
        if response_id not in RESPONSE_CACHE:
            await callback.answer("❌ Файл устарел или не найден", show_alert=True)
            return
            
        content = RESPONSE_CACHE[response_id]
        
        if format_type == "txt":
            file_bytes = content.encode('utf-8')
            file = BufferedInputFile(file_bytes, filename="response.txt")
            await callback.message.reply_document(
                document=file,
                caption="📄 Ваш ответ в формате TXT"
            )
            # Удаляем кнопки после нажатия
            await callback.message.edit_reply_markup(reply_markup=None)
            await callback.answer()
            
    except Exception as e:
        logger.error(f"Error downloading file: {e}")
        await callback.answer("❌ Ошибка при формировании файла", show_alert=True)


# ============== ОБРАБОТЧИКИ КОНТЕНТА ==============

import re
from aiogram.enums import ParseMode


def convert_markdown_to_html(text: str) -> str:
    """
    Конвертирует Markdown в HTML для Telegram.
    Поддерживает: жирный, курсив, код, блоки кода.
    """
    # Сначала обрабатываем блоки кода (чтобы не затронуть их содержимое)
    code_blocks = []
    
    def save_code_block(match):
        lang = match.group(1) or ""
        code = match.group(2)
        # Экранируем HTML в коде
        code = code.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if lang:
            code_blocks.append(f'<pre><code class="language-{lang}">{code}</code></pre>')
        else:
            code_blocks.append(f'<pre><code>{code}</code></pre>')
        return f"<<<CODE_BLOCK_{len(code_blocks) - 1}>>>"
    
    # Блоки кода ```language\ncode```
    text = re.sub(r'```(\w*)\n?(.*?)```', save_code_block, text, flags=re.DOTALL)
    
    # Инлайн код `code`
    inline_codes = []
    def save_inline_code(match):
        code = match.group(1)
        code = code.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        inline_codes.append(f'<code>{code}</code>')
        return f"<<<INLINE_CODE_{len(inline_codes) - 1}>>>"
    
    text = re.sub(r'`([^`]+)`', save_inline_code, text)
    
    # Экранируем HTML в остальном тексте
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    
    # Жирный текст **text** или *text*
    text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'(?<![*\w])\*([^*\n]+)\*(?![*\w])', r'<b>\1</b>', text)
    
    # Курсив _text_
    text = re.sub(r'(?<![_\w])_([^_\n]+)_(?![_\w])', r'<i>\1</i>', text)
    
    # Зачёркнутый ~text~
    text = re.sub(r'~([^~]+)~', r'<s>\1</s>', text)
    
    # Спойлер ||text||
    text = re.sub(r'\|\|([^|]+)\|\|', r'<tg-spoiler>\1</tg-spoiler>', text)
    
    # Подчёркнутый __text__
    text = re.sub(r'__([^_]+)__', r'<u>\1</u>', text)
    
    # Восстанавливаем блоки кода
    for i, block in enumerate(code_blocks):
        text = text.replace(f"&lt;&lt;&lt;CODE_BLOCK_{i}&gt;&gt;&gt;", block)
    
    for i, code in enumerate(inline_codes):
        text = text.replace(f"&lt;&lt;&lt;INLINE_CODE_{i}&gt;&gt;&gt;", code)
    
    return text


def clean_markdown(text: str) -> str:
    """Удаляет символы markdown из текста"""
    # Жирный
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Курсив
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Код
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'```[\w]*\n(.+?)```', r'\1', text, flags=re.DOTALL)
    # Заголовки
    text = re.sub(r'^#+\s+(.+)$', r'\1', text, flags=re.MULTILINE)
    # Списки
    # text = re.sub(r'^\s*[\-\*]\s+', '', text, flags=re.MULTILINE) # Оставляем списки, они полезны
    return text


async def send_response(message: Message, response: str) -> None:
    """Отправляет ответ с HTML форматированием, если слишком длинный — в файле"""
    if len(response) <= MAX_TELEGRAM_MESSAGE_LENGTH:
        try:
            # Пробуем отправить с HTML форматированием
            html_response = convert_markdown_to_html(response)
            await message.reply(html_response, parse_mode=ParseMode.HTML)
        except Exception as e:
            # Если форматирование сломалось, отправляем без него
            logger.warning(f"HTML parse error: {e}")
            try:
                await message.reply(response)
            except Exception:
                # Если и так не получилось, отправляем DOCX + кнопку на TXT
                clean_text = clean_markdown(response)
                response_id = str(uuid.uuid4())
                RESPONSE_CACHE[response_id] = clean_text
                
                try:
                    # Создаем DOCX
                    doc = Document()
                    doc.add_paragraph(clean_text)
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    file = BufferedInputFile(buffer.read(), filename="response.docx")
                    
                    await message.reply_document(
                        document=file,
                        caption="📄 **Не удалось отправить сообщение**\n"
                                "Отправляю в формате DOCX (без форматирования).",
                        reply_markup=get_txt_download_keyboard(response_id),
                        parse_mode="Markdown"
                    )
                except Exception as ex:
                    # Если DOCX не создался, шлем TXT
                    logger.error(f"DOCX gen error: {ex}")
                    file_bytes = clean_text.encode('utf-8')
                    file = BufferedInputFile(file_bytes, filename="response.txt")
                    await message.reply_document(
                        document=file,
                        caption="📄 Отправляю файлом."
                    )

    else:
        # Ответ слишком длинный - отправляем DOCX + кнопку на TXT
        clean_text = clean_markdown(response)
        response_id = str(uuid.uuid4())
        RESPONSE_CACHE[response_id] = clean_text
        
        try:
            # Создаем DOCX
            doc = Document()
            doc.add_paragraph(clean_text)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            file = BufferedInputFile(buffer.read(), filename="response.docx")
            
            await message.reply_document(
                document=file,
                caption="📄 **Ответ слишком длинный**\n"
                        "Отправляю в формате DOCX.",
                reply_markup=get_txt_download_keyboard(response_id),
                parse_mode="Markdown"
            )
        except Exception as ex:
             # Если DOCX не создался, шлем TXT
            logger.error(f"DOCX gen error: {ex}")
            file_bytes = clean_text.encode('utf-8')
            file = BufferedInputFile(file_bytes, filename="response.txt")
            await message.reply_document(
                document=file,
                caption="📄 Ответ слишком длинный, отправляю файлом."
            )


@router.message(F.photo)
async def handle_photo(message: Message, bot: Bot) -> None:
    """Обработчик фотографий"""
    user_id = message.from_user.id
    
    # Получаем фото максимального размера
    photo = message.photo[-1]
    
    # Показываем индикатор
    await bot.send_chat_action(message.chat.id, "typing")
    
    # Скачиваем фото
    file = await bot.get_file(photo.file_id)
    file_data = await bot.download_file(file.file_path)
    image_bytes = file_data.read()
    
    # Проверяем режим редактирования
    if conversation_manager.is_edit_mode(user_id):
        # Сохраняем изображение для редактирования
        conversation_manager.set_user_image(user_id, image_bytes)
        
        # Если есть подпись - сразу редактируем
        if message.caption:
            await bot.send_chat_action(message.chat.id, "upload_photo")
            status_msg = await message.reply("🎨 Редактирую изображение...")
            
            result_url, result_text = await edit_image_with_dalle(image_bytes, message.caption)
            
            if result_url:
                try:
                    # Если это data URL, декодируем
                    if result_url.startswith("data:"):
                        import base64
                        b64_data = result_url.split(",")[1]
                        edited_bytes = base64.b64decode(b64_data)
                    else:
                        # Скачиваем по URL
                        import aiohttp
                        async with aiohttp.ClientSession() as session:
                            async with session.get(result_url, timeout=aiohttp.ClientTimeout(total=60)) as resp:
                                edited_bytes = await resp.read()
                    
                    # Сохраняем отредактированное изображение как новое
                    conversation_manager.set_user_image(user_id, edited_bytes)
                    
                    result_photo = BufferedInputFile(edited_bytes, filename="edited.png")
                    await status_msg.delete()
                    await message.reply_photo(
                        photo=result_photo,
                        caption=f"✅ Готово! Отредактировано по запросу:\n{message.caption[:200]}"
                    )
                except Exception as e:
                    logger.error(f"Error sending edited image: {e}")
                    await status_msg.edit_text(f"❌ Ошибка: {str(e)[:100]}")
            else:
                await status_msg.edit_text(result_text or "❌ Ошибка редактирования")
        else:
            await message.reply(
                "✅ Изображение сохранено!\n\n"
                "Теперь напиши что нужно изменить, например:\n"
                "• добавь солнечные очки\n"
                "• сделай фон синим\n"
                "• добавь шляпу"
            )
        return
    
    # Обычный режим - анализ изображения
    image_base64 = encode_image_to_base64(image_bytes)
    
    # Определяем MIME-тип
    mime_type = "image/jpeg"
    if file.file_path and file.file_path.lower().endswith('.png'):
        mime_type = "image/png"
    
    # Получаем подпись (если есть)
    user_text = message.caption or "Опиши это изображение подробно."
    
    # Добавляем в историю
    conversation_manager.add_message(user_id, "user", f"[Изображение] {user_text}", MAX_HISTORY_MESSAGES)
    
    # Получаем историю и модель
    messages = conversation_manager.get_messages_for_api(user_id, SYSTEM_PROMPT)
    model = conversation_manager.get_user_model(user_id)
    
    # Получаем ответ (для изображений используется vision модель)
    response = await get_chat_response(messages, model=model, image_base64=image_base64, image_mime_type=mime_type)
    
    # Сохраняем ответ
    conversation_manager.add_message(user_id, "assistant", response, MAX_HISTORY_MESSAGES)
    
    # Отправляем
    await send_response(message, response)


@router.message(F.document)
async def handle_document(message: Message, bot: Bot) -> None:
    """Обработчик документов (PDF, DOCX, TXT)"""
    try:
        user_id = message.from_user.id
        document = message.document
        file_name = document.file_name or "document"
        
        # Проверяем формат файла
        supported_formats = ('.pdf', '.docx', '.txt')
        if not any(file_name.lower().endswith(ext) for ext in supported_formats):
            await message.reply(
                "⚠️ Поддерживаются форматы: PDF, DOCX, TXT\n"
                "Отправь документ в одном из этих форматов."
            )
            return
        
        # Показываем индикатор загрузки
        status_msg = await message.reply("📥 Загружаю файл...")
        
        try:
            # Скачиваем документ
            file = await bot.get_file(document.file_id)
            file_data = await bot.download_file(file.file_path)
            file_bytes = file_data.read()
        except Exception as e:
            logger.error(f"Error downloading document: {e}")
            await status_msg.edit_text(f"❌ Ошибка загрузки файла: {str(e)[:100]}")
            return
        
        # === РЕЖИМ ШАБЛОНОВ ===
        if conversation_manager.is_template_mode(user_id):
            if not file_name.lower().endswith('.docx'):
                await status_msg.edit_text(
                    "❌ В режиме шаблонов поддерживается только **DOCX**!\n\n"
                    "Отправь документ Word (.docx)",
                    parse_mode="Markdown"
                )
                return
            
            # Сохраняем как шаблон
            conversation_manager.set_template_doc(user_id, file_bytes, file_name)
            
            # Получаем структуру документа для показа пользователю
            doc_structure = await get_docx_structure_for_ai(file_bytes)
            
            # Обрезаем если слишком длинный
            if len(doc_structure) > 2000:
                doc_structure = doc_structure[:2000] + "\n\n[... документ обрезан для превью ...]"
            
            await status_msg.edit_text(
                f"✅ **Шаблон загружен:** `{file_name}`\n\n"
                f"📄 **Содержимое:**\n```\n{doc_structure[:1500]}\n```\n\n"
                "🔧 **Теперь опиши что нужно заменить:**\n"
                "• _\"Замени [старый текст] на [новый текст]\"_\n"
                "• _\"Измени ООО Ромашка на ООО Василёк\"_\n"
                "• _\"Поменяй дату 01.01.2025 на 15.03.2026\"_",
                parse_mode="Markdown"
            )
            return
        
        # === ОБЫЧНЫЙ РЕЖИМ - анализ документа ===
        await status_msg.edit_text("⚙️ Обрабатываю файл...")
        
        # Извлекаем текст
        extracted_text = await extract_text_from_file(file_bytes, file_name)
        
        if not extracted_text or extracted_text.startswith("Ошибка"):
            error_text = extracted_text or "Не удалось извлечь текст"
            await status_msg.edit_text(f"❌ {error_text}")
            return
        
        # Обрезаем текст, если он слишком длинный (увеличили лимит)
        max_chars = 100000
        if len(extracted_text) > max_chars:
            extracted_text = extracted_text[:max_chars] + "\n\n[... текст обрезан ...]"
        
        # Добавляем в историю КАК КЕЙС (не как сообщение пользователя, чтобы бот не отвечал сам себе)
        # Но мы хотим, чтобы бот знал контекст. 
        # Сохраняем это как сообщение пользователя с пометкой
        conversation_manager.add_message(user_id, "user", f"[Документ: {file_name}]\n{extracted_text}", MAX_HISTORY_MESSAGES)
        
        # Проверяем, был ли вопрос (caption)
        user_question = message.caption
        
        if user_question:
            # Если есть вопрос - отвечаем на него
            await status_msg.edit_text("Ищу ответ на Ваш вопрос...")
            await bot.send_chat_action(message.chat.id, "typing")
            
            # Добавляем вопрос пользователя в историю ОТДЕЛЬНО
            conversation_manager.add_message(user_id, "user", user_question, MAX_HISTORY_MESSAGES)
            
            # Получаем историю и модель
            messages = conversation_manager.get_messages_for_api(user_id, SYSTEM_PROMPT)
            model = conversation_manager.get_user_model(user_id)
            
            # Получаем ответ
            response = await get_chat_response(messages, model=model)
            
            # Сохраняем ответ
            conversation_manager.add_message(user_id, "assistant", response, MAX_HISTORY_MESSAGES)
            
            # Редактируем сообщение с ответом
            await send_response_edit(status_msg, message, response)
        else:
            # Если вопроса нет - подтверждаем реакцией
            await status_msg.delete()
            await add_heart_reaction(message, bot)

    except Exception as e:
        logger.error(f"Unhandled error in handle_document: {e}")
        try:
            await message.reply(f"❌ Критическая ошибка при обработке файла: {e}")
        except Exception:
            pass


async def send_response_edit(status_msg: Message, original_msg: Message, response: str) -> None:
    """Редактирует статусное сообщение с финальным ответом"""
    from aiogram.enums import ParseMode
    
    if len(response) <= MAX_TELEGRAM_MESSAGE_LENGTH:
        try:
            html_response = convert_markdown_to_html(response)
            await status_msg.edit_text(html_response, parse_mode=ParseMode.HTML)
        except Exception as e:
            logger.warning(f"HTML edit error: {e}")
            try:
                await status_msg.edit_text(response)
            except Exception:
                # Если редактирование не сработало
                await status_msg.delete()
                
                clean_text = clean_markdown(response)
                response_id = str(uuid.uuid4())
                RESPONSE_CACHE[response_id] = clean_text
                
                try:
                    # Создаем DOCX
                    doc = Document()
                    doc.add_paragraph(clean_text)
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    file = BufferedInputFile(buffer.read(), filename="response.docx")
                    
                    await original_msg.reply_document(
                        document=file,
                        caption="📄 **Не удалось отправить сообщение**\n"
                                "Отправляю в формате DOCX.",
                        reply_markup=get_txt_download_keyboard(response_id),
                        parse_mode="Markdown"
                    )
                except Exception:
                     file_bytes = clean_text.encode('utf-8')
                     file = BufferedInputFile(file_bytes, filename="response.txt")
                     await original_msg.reply_document(
                        document=file,
                        caption="📄 Отправляю файлом."
                     )
    else:
        # Ответ слишком длинный
        await status_msg.delete()
        
        clean_text = clean_markdown(response)
        response_id = str(uuid.uuid4())
        RESPONSE_CACHE[response_id] = clean_text
        
        try:
            # Создаем DOCX
            doc = Document()
            doc.add_paragraph(clean_text)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            file = BufferedInputFile(buffer.read(), filename="response.docx")
            
            await original_msg.reply_document(
                document=file,
                caption="📄 **Ответ слишком длинный**\n"
                        "Отправляю в формате DOCX.",
                reply_markup=get_txt_download_keyboard(response_id),
                parse_mode="Markdown"
            )
        except Exception:
             file_bytes = clean_text.encode('utf-8')
             file = BufferedInputFile(file_bytes, filename="response.txt")
             await original_msg.reply_document(
                document=file,
                caption="📄 Ответ слишком длинный, отправляю файлом."
             )


@router.message(F.voice)
async def handle_voice(message: Message, bot: Bot) -> None:
    """Обработчик голосовых сообщений"""
    try:
        user_id = message.from_user.id
        
        # Статус загрузки
        status_msg = await message.reply("🎤 Загружаю голосовое сообщение...")
        
        try:
            # Скачиваем голосовое
            file = await bot.get_file(message.voice.file_id)
            file_data = await bot.download_file(file.file_path)
            audio_bytes = file_data.read()
        except Exception as e:
            logger.error(f"Error downloading voice: {e}")
            await status_msg.edit_text(f"❌ Ошибка загрузки голосового: {str(e)[:100]}")
            return
        
        # Транскрибируем
        await status_msg.edit_text("🔊 Распознаю речь...")
        
        transcribed_text = await transcribe_audio(audio_bytes, "ogg")
        
        if transcribed_text.startswith("❌"):
            await status_msg.edit_text(transcribed_text)
            return
        
        # Показываем распознанный текст
        await status_msg.edit_text(f"📝 Распознано: _{transcribed_text}_\n\nИщу ответ на Ваш вопрос...", parse_mode="Markdown")
        
        await bot.send_chat_action(message.chat.id, "typing")
        
        # Добавляем в историю
        conversation_manager.add_message(user_id, "user", transcribed_text, MAX_HISTORY_MESSAGES)
        
        # Получаем историю и модель
        messages = conversation_manager.get_messages_for_api(user_id, SYSTEM_PROMPT)
        model = conversation_manager.get_user_model(user_id)
        
        # Получаем ответ
        response = await get_chat_response(messages, model=model)
        
        # Сохраняем ответ
        conversation_manager.add_message(user_id, "assistant", response, MAX_HISTORY_MESSAGES)
        
        # Редактируем сообщение с ответом
        await send_response_edit(status_msg, message, response)

    except Exception as e:
        logger.error(f"Unhandled error in handle_voice: {e}")
        try:
            await message.reply(f"❌ Критическая ошибка при обработке голосового: {e}")
        except Exception:
            pass


@router.message(F.text)
@safe_handler
async def handle_text(message: Message, bot: Bot, **kwargs) -> None:
    """Обработчик текстовых сообщений"""
    user_id = message.from_user.id
    user_text = message.text
    
    logger.info(f"Received text message from {user_id}: {user_text[:50]}...")
    
    # Игнорируем пустые сообщения
    if not user_text or not user_text.strip():
        return
    
    # Проверяем режим редактирования
    if conversation_manager.is_edit_mode(user_id):
        saved_image = conversation_manager.get_user_image(user_id)
        
        if saved_image:
            # Редактируем сохранённое изображение
            await bot.send_chat_action(message.chat.id, "upload_photo")
            status_msg = await message.reply("🎨 Редактирую изображение...")
            
            result_url, result_text = await edit_image_with_dalle(saved_image, user_text)
            
            if result_url:
                try:
                    # Если это data URL, декодируем
                    if result_url.startswith("data:"):
                        import base64
                        b64_data = result_url.split(",")[1]
                        edited_bytes = base64.b64decode(b64_data)
                    else:
                        # Скачиваем по URL
                        import aiohttp
                        async with aiohttp.ClientSession() as session:
                            async with session.get(result_url, timeout=aiohttp.ClientTimeout(total=60)) as resp:
                                edited_bytes = await resp.read()
                    
                    # Сохраняем отредактированное изображение как новое
                    conversation_manager.set_user_image(user_id, edited_bytes)
                    
                    result_photo = BufferedInputFile(edited_bytes, filename="edited.png")
                    await status_msg.delete()
                    await message.reply_photo(
                        photo=result_photo,
                        caption=f"✅ Готово: {user_text[:200]}\n\n💡 Можешь продолжить редактировать или отправить новое фото."
                    )
                except Exception as e:
                    logger.error(f"Error sending edited image: {e}")
                    try:
                        await status_msg.edit_text(f"❌ Ошибка: {str(e)[:100]}")
                    except Exception:
                        await message.reply(f"❌ Ошибка: {str(e)[:100]}")
            else:
                try:
                    await status_msg.edit_text(result_text or "❌ Ошибка редактирования")
                except Exception:
                    await message.reply(result_text or "❌ Ошибка редактирования")
            return
        else:
            await message.reply(
                "📸 Сначала отправь фото для редактирования!\n\n"
                "Отправь изображение, которое хочешь редактировать."
            )
            return
    
    # Проверяем DALL-E режим
    if conversation_manager.is_dalle_mode(user_id):
        await bot.send_chat_action(message.chat.id, "upload_photo")
        
        # Проверяем, есть ли сохранённое изображение
        saved_dalle_image = conversation_manager.get_dalle_image(user_id)
        
        if saved_dalle_image:
            # Редактируем существующее изображение
            status_msg = await message.reply("🎨 Редактирую изображение...")
            
            result_url, result_text = await edit_image_with_dalle(saved_dalle_image, user_text)
            
            if result_url:
                try:
                    if result_url.startswith("data:"):
                        import base64
                        b64_data = result_url.split(",")[1]
                        edited_bytes = base64.b64decode(b64_data)
                    else:
                        import aiohttp
                        async with aiohttp.ClientSession() as session:
                            async with session.get(result_url, timeout=aiohttp.ClientTimeout(total=60)) as resp:
                                edited_bytes = await resp.read()
                    
                    # Сохраняем новое изображение
                    conversation_manager.set_dalle_image(user_id, edited_bytes)
                    
                    result_photo = BufferedInputFile(edited_bytes, filename="dalle_edited.png")
                    await status_msg.delete()
                    await message.reply_photo(
                        photo=result_photo,
                        caption=f"✅ {user_text[:200]}\n\n💡 Продолжай редактировать или напиши новый промпт для новой картинки"
                    )
                except Exception as e:
                    logger.error(f"Error sending edited DALL-E image: {e}")
                    await status_msg.edit_text(f"❌ Ошибка: {str(e)[:100]}")
            else:
                await status_msg.edit_text(result_text or "❌ Ошибка редактирования")
        else:
            # Генерируем новое изображение
            status_msg = await message.reply("🖼 Генерирую изображение...")
            
            image_url, result = await generate_image(user_text)
            
            if image_url:
                try:
                    import aiohttp
                    async with aiohttp.ClientSession() as session:
                        async with session.get(image_url, timeout=aiohttp.ClientTimeout(total=60)) as resp:
                            if resp.status == 200:
                                image_data = await resp.read()
                            else:
                                raise Exception(f"HTTP {resp.status}")
                    
                    # Сохраняем для последующего редактирования
                    conversation_manager.set_dalle_image(user_id, image_data)
                    
                    photo = BufferedInputFile(image_data, filename="dalle_generated.png")
                    
                    caption = f"🖼 {user_text[:150]}"
                    if result and result != user_text:
                        caption += f"\n\n📝 Промпт DALL-E: {result[:150]}"
                    caption += "\n\n💡 Напиши что изменить, чтобы отредактировать"
                    
                    await status_msg.delete()
                    await message.reply_photo(photo=photo, caption=caption)
                    
                except Exception as e:
                    logger.error(f"Error sending DALL-E image: {e}")
                    await status_msg.edit_text(f"❌ Ошибка: {str(e)[:100]}")
            else:
                await status_msg.edit_text(result or "❌ Не удалось сгенерировать")
        return
    
    # Проверяем режим шаблонов
    if conversation_manager.is_template_mode(user_id):
        template_doc = conversation_manager.get_template_doc(user_id)
        template_name = conversation_manager.get_template_name(user_id)
        
        if not template_doc:
            await message.reply(
                "📄 Сначала отправь DOCX документ (шаблон)!\n\n"
                "Загрузи файл, который хочешь редактировать."
            )
            return
        
        # Получаем структуру документа для AI
        status_msg = await message.reply("🔍 Анализирую документ и готовлю замены...")
        
        await bot.send_chat_action(message.chat.id, "typing")
        
        doc_structure = await get_docx_structure_for_ai(template_doc)
        if len(doc_structure) > 10000:
            doc_structure = doc_structure[:10000] + "\n[...обрезано...]"
        
        # Промпт для AI чтобы он распарсил замены - УМНЫЙ режим
        ai_prompt = f"""Ты — умный помощник для редактирования документов. Твоя задача — понять, что хочет пользователь, даже если он выражается неточно или примерно.

=== ДОКУМЕНТ ===
{doc_structure}

=== ЗАПРОС ПОЛЬЗОВАТЕЛЯ ===
{user_text}

=== ТВОЯ ЗАДАЧА ===
Проанализируй документ и пойми, что именно пользователь хочет изменить. Пользователь может:
- Говорить приблизительно ("поменяй название фирмы на Ромашка" — найди ВСЕ названия компаний в документе)
- Указывать частично ("замени дату" — найди даты в документе и замени на указанную)
- Описывать суть ("сделай договор на другую компанию ООО Тест" — замени ВСЕ упоминания старой компании)
- Говорить про тип данных ("поменяй телефон на +7999..." — найди телефоны в документе)
- Просить изменить реквизиты ("ИНН замени на 1234567890" — найди ИНН и замени)
- Говорить про ФИО ("ФИО директора замени на Иванов И.И." — найди ФИО в соответствующем контексте)
- Говорить про должности и подписи ("подпись преподавателя", "ФИО исполнителя" и т.д.)

АЛГОРИТМ:
1. Пойми НАМЕРЕНИЕ пользователя — что он хочет изменить по смыслу
2. Найди в документе ВСЕ соответствующие фрагменты (даже если пользователь не назвал их точно)
3. Ищи по КОНТЕКСТУ: если пользователь говорит "ФИО преподавателя" — найди в документе где упоминается преподаватель и его ФИО рядом
4. Создай словарь замен, где ключи — ТОЧНЫЕ строки из документа

ВЕРНИ JSON в формате:
{{"точный_текст_из_документа_1": "новое_значение_1", "точный_текст_из_документа_2": "новое_значение_2"}}

ПРИМЕРЫ:
- Пользователь: "компанию поменяй на ООО Василёк"
  Документ содержит: "ООО Ромашка", "ООО «Ромашка»"  
  Ответ: {{"ООО Ромашка": "ООО Василёк", "ООО «Ромашка»": "ООО «Василёк»"}}

- Пользователь: "дату на 15 марта 2026"
  Документ содержит: "01 января 2025 г.", "01.01.2025"
  Ответ: {{"01 января 2025 г.": "15 марта 2026 г.", "01.01.2025": "15.03.2026"}}

- Пользователь: "сумму сделай 500 тысяч"
  Документ содержит: "100 000 (Сто тысяч) рублей"
  Ответ: {{"100 000 (Сто тысяч) рублей": "500 000 (Пятьсот тысяч) рублей"}}

- Пользователь: "ФИО преподавателя замени на Дагаев А.В."
  Документ содержит: "Преподаватель: Иванов П.С.", "Иванов Пётр Сергеевич"
  Ответ: {{"Иванов П.С.": "Дагаев А.В.", "Иванов Пётр Сергеевич": "Дагаев А.В."}}

- Пользователь: "ФИО директора на Петров"
  Документ содержит: "Директор ____________ Сидоров А.А."
  Ответ: {{"Сидоров А.А.": "Петров А.А."}}

- Пользователь: "поменяй студента на Козлов"
  Документ содержит: "Студент группы ИТ-21: Смирнов Алексей Игоревич"
  Ответ: {{"Смирнов Алексей Игоревич": "Козлов Алексей Игоревич"}}

КРИТИЧЕСКИ ВАЖНО:
- Ключи словаря должны быть ТОЧНЫМИ копиями текста из документа (с пробелами, скобками и т.д.)
- Отвечай ТОЛЬКО валидным JSON, без пояснений и без markdown
- Если пользователь говорит расплывчато — используй контекст документа чтобы понять что менять  
- Ищи ФИО рядом с указанной ролью (преподаватель, студент, директор и т.д.)
- Если совсем непонятно — верни {{"_error": "Уточни, что именно заменить"}}

JSON:"""

        # Получаем замены от AI
        model = conversation_manager.get_user_model(user_id)
        ai_response = await get_chat_response([
            {"role": "system", "content": "Ты эксперт по редактированию документов. Ты умеешь понимать неточные запросы пользователя и находить в документе нужные фрагменты для замены. Отвечай ТОЛЬКО валидным JSON."},
            {"role": "user", "content": ai_prompt}
        ], model=model)
        
        # Логируем ответ AI для отладки
        logger.info(f"Template AI response: {ai_response[:500]}")
        
        # Парсим JSON - улучшенная версия
        import json
        import re
        
        replacements = None
        
        try:
            # Способ 1: Напрямую как JSON
            replacements = json.loads(ai_response.strip())
        except json.JSONDecodeError:
            pass
        
        if replacements is None:
            try:
                # Способ 2: Ищем JSON между ``` блоками
                code_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', ai_response, re.DOTALL)
                if code_match:
                    replacements = json.loads(code_match.group(1))
            except json.JSONDecodeError:
                pass
        
        if replacements is None:
            try:
                # Способ 3: Ищем любой JSON объект (с вложенностью)
                # Находим первую { и последнюю }
                start_idx = ai_response.find('{')
                end_idx = ai_response.rfind('}')
                if start_idx != -1 and end_idx > start_idx:
                    json_str = ai_response[start_idx:end_idx + 1]
                    replacements = json.loads(json_str)
            except json.JSONDecodeError:
                pass
        
        if replacements is None:
            logger.error(f"Failed to parse JSON from AI response: {ai_response}")
            await status_msg.edit_text(
                "🤔 Не смог разобрать ответ. Попробуй переформулировать:\n\n"
                f"Твой запрос: _{user_text}_\n\n"
                "Примеры:\n"
                "• _\"ФИО преподавателя замени на Иванов И.И.\"_\n"
                "• _\"Название компании поменяй на ООО Тест\"_\n"
                "• _\"Дату сделай 15.03.2026\"_",
                parse_mode="Markdown"
            )
            return
        
        # Проверяем на ошибку от AI
        if "_error" in replacements:
            await status_msg.edit_text(
                f"🤔 {replacements['_error']}\n\n"
                "Попробуй сформулировать иначе, например:\n"
                "• _\"Название компании поменяй на ООО Ромашка\"_\n"
                "• _\"Дату сделай 15.03.2026\"_\n"
                "• _\"Сумму измени на 500 000 рублей\"_",
                parse_mode="Markdown"
            )
            return
        
        if not replacements:
            await status_msg.edit_text(
                "🤔 Не понял, что нужно заменить.\n\n"
                "Опиши подробнее, например:\n"
                "• _\"Поменяй компанию на ООО Тест\"_\n"
                "• _\"Дату договора сделай 15 марта 2026\"_\n"
                "• _\"Телефон замени на +7 999 123-45-67\"_",
                parse_mode="Markdown"
            )
            return
        
        # Показываем что будем менять
        await status_msg.edit_text("📝 Применяю изменения к документу...")
        
        # Редактируем документ
        try:
            edited_doc = await edit_docx_with_replacements(template_doc, replacements)
            
            # Формируем имя файла
            base_name = template_name.rsplit('.', 1)[0] if template_name else "document"
            new_filename = f"{base_name}_edited.docx"
            
            # Отправляем результат
            doc_file = BufferedInputFile(edited_doc, filename=new_filename)
            
            # Формируем список замен для caption
            replacements_text = "\n".join([f"• `{old}` → `{new}`" for old, new in list(replacements.items())[:5]])
            if len(replacements) > 5:
                replacements_text += f"\n... и ещё {len(replacements) - 5} замен"
            
            await status_msg.delete()
            await message.reply_document(
                document=doc_file,
                caption=f"✅ **Документ отредактирован!**\n\n"
                        f"📝 **Замены:**\n{replacements_text}\n\n"
                        f"💡 Можешь загрузить новый шаблон или продолжить редактировать текущий.",
                parse_mode="Markdown"
            )
            
            # Сохраняем отредактированный как новый шаблон
            conversation_manager.set_template_doc(user_id, edited_doc, new_filename)
            
        except Exception as e:
            logger.error(f"Error editing document: {e}")
            await status_msg.edit_text(f"❌ Ошибка при редактировании: {str(e)[:100]}")
        
        return
    
    # Обычный текстовый запрос - показываем анимированный статус
    status_msg = await message.reply("Ищу ответ на Ваш вопрос...")
    
    logger.info(f"Processing normal text for {user_id}, sending typing action")
    await bot.send_chat_action(message.chat.id, "typing")
    
    # Добавляем сообщение в историю
    conversation_manager.add_message(user_id, "user", user_text, MAX_HISTORY_MESSAGES)
    
    # Определяем системный промпт (стандартный + кастомный если есть)
    custom_prompt = conversation_manager.get_active_custom_prompt(user_id)
    if custom_prompt:
        system_prompt = f"{SYSTEM_PROMPT}\n\nДополнительные инструкции пользователя:\n{custom_prompt}"
    else:
        system_prompt = SYSTEM_PROMPT
    
    # Получаем историю и модель пользователя
    messages = conversation_manager.get_messages_for_api(user_id, system_prompt)
    model = conversation_manager.get_user_model(user_id)
    
    # Получаем ответ
    logger.info(f"Sending request to OpenAI (model={model})...")
    response = await get_chat_response(messages, model=model)
    logger.info(f"Received response from OpenAI: {len(response)} chars")
    
    # Добавляем реакцию сердечком на вопрос пользователя
    await add_heart_reaction(message, bot)
    
    # Сохраняем ответ в историю
    conversation_manager.add_message(user_id, "assistant", response, MAX_HISTORY_MESSAGES)
    
    # Редактируем статусное сообщение с ответом
    await send_response_edit(status_msg, message, response)


